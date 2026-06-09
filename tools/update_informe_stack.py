from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


DOC_PATH = Path(r"D:/FDM/Alonso_Figueroa_informe_actualizado.docx")


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def replace_first_containing(doc, fragment, replacement):
    for paragraph in doc.paragraphs:
        if fragment in paragraph.text:
            set_paragraph_text(paragraph, replacement)
            return
    raise ValueError(f"Paragraph fragment not found: {fragment!r}")


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_p.getparent().replace(new_p, new_paragraph._p)
    new_paragraph.style = paragraph.style
    if paragraph._p.pPr is not None:
        new_paragraph._p.insert(0, deepcopy(paragraph._p.pPr))
    new_paragraph.add_run(text)
    return new_paragraph


def update_references(doc):
    replace_first_containing(
        doc,
        "Express.js. (s.f.). Express - Node.js web application framework.",
        "Astro. (s.f.). Astro Documentation. Recuperado de https://docs.astro.build/",
    )
    replace_first_containing(
        doc,
        "PostgreSQL Global Development Group. (s.f.). PostgreSQL Documentation.",
        "NestJS. (s.f.). Documentation. Recuperado de https://docs.nestjs.com/",
    )
    replace_first_containing(
        doc,
        "Mozilla Developer Network. (s.f.). Responsive design.",
        "Netlify. (s.f.). Deploy overview. Recuperado de https://docs.netlify.com/deploy/deploy-overview/",
    )

    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Netlify. (s.f.). Deploy overview."):
            anchor = paragraph
            break
    if anchor is None:
        raise ValueError("Reference anchor not found.")

    anchor = insert_paragraph_after(
        anchor,
        "Supabase. (s.f.). Database. Recuperado de https://supabase.com/docs/guides/database/overview",
    )
    anchor = insert_paragraph_after(
        anchor,
        "Supabase. (s.f.). Storage. Recuperado de https://supabase.com/docs/guides/storage",
    )
    anchor = insert_paragraph_after(
        anchor,
        "TypeORM. (s.f.). TypeORM Documentation. Recuperado de https://typeorm.io/docs/getting-started",
    )
    insert_paragraph_after(
        anchor,
        "Mozilla Developer Network. (s.f.). Responsive design. Recuperado de https://developer.mozilla.org/",
    )


def update_anexo(doc):
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("• La empresa solicitante cuenta o contará con acceso técnico"):
            anchor = paragraph
            break
    if anchor is None:
        raise ValueError("Anexo anchor not found.")

    anchor = insert_paragraph_after(
        anchor,
        "• El despliegue web se realizará en Netlify, usando variables de entorno para separar credenciales y configuración de producción.",
    )
    insert_paragraph_after(
        anchor,
        "• La base de datos y el almacenamiento documental se administrarán en Supabase, con PostgreSQL para datos relacionales y Storage para PDFs originales, firmados y recursos asociados.",
    )


def main():
    doc = Document(str(DOC_PATH))

    replacements = {
        "La propuesta consiste en desarrollar un sistema web responsive que actúe como una capa de gestión": (
            "La propuesta consiste en desarrollar un sistema web responsive que actúe como una capa de gestión entre el usuario final y el proveedor de firma electrónica avanzada Firma.cl. El stack actualizado considera un frontend Astro con componentes React, una API backend NestJS sobre Node.js, despliegue web en Netlify y Supabase como servicio gestionado para base de datos PostgreSQL y almacenamiento documental. En este enfoque, ClaveÚnica no se plantea como el objetivo central del sistema, sino como un mecanismo de validación de identidad que puede formar parte del flujo habilitado por el proveedor. La finalidad es simplificar la experiencia del usuario, ordenar la operación interna y mantener evidencia de las acciones realizadas durante el ciclo de vida del documento."
        ),
        "La relevancia del proyecto se relaciona con la transformación digital de procesos documentales": (
            "La relevancia del proyecto se relaciona con la transformación digital de procesos documentales en organizaciones que requieren seguridad, control y eficiencia. La firma electrónica avanzada permite reducir desplazamientos, disminuir tiempos de operación y fortalecer la trazabilidad de documentos relevantes. Además, el proyecto se ajusta al área de especialidad, ya que integra desarrollo web con Astro, React y Tailwind/CSS responsive, backend NestJS, base de datos Supabase PostgreSQL, almacenamiento Supabase Storage, seguridad, manejo de archivos, integración con servicios externos, despliegue en Netlify y análisis de requerimientos."
        ),
        "Desde el punto de vista técnico y formativo, la propuesta permite aplicar competencias": (
            "Desde el punto de vista técnico y formativo, la propuesta permite aplicar competencias propias del área de informática, tales como desarrollo web, modelado de datos, integración con servicios externos, autenticación, control de roles, validación de archivos, manejo de documentos PDF, despliegue en Netlify y diseño de interfaces responsive mediante Tailwind/CSS propio. Por ello, no se limita a un sitio informativo, sino que propone un sistema funcional con flujo de negocio, reglas, formularios, almacenamiento de información en Supabase PostgreSQL y resguardo de archivos en Supabase Storage."
        ),
        "La solución propuesta es el desarrollo de un sistema web responsive para la gestión": (
            "La solución propuesta es el desarrollo de un sistema web responsive para la gestión, posicionamiento visual y firma electrónica avanzada de documentos PDF. El sistema permitirá a usuarios autorizados ingresar a la plataforma, cargar documentos, seleccionar la página y ubicación donde se mostrará la firma, enviar la solicitud al servicio de integración con Firma.cl, recibir el resultado firmado, registrar su trazabilidad en Supabase PostgreSQL y almacenarlo en Supabase Storage para su posterior descarga controlada."
        ),
        "La arquitectura considera una interfaz de usuario, un backend de servicios y una base de datos relacional.": (
            "La arquitectura considera una interfaz de usuario, un backend de servicios, una base de datos relacional gestionada y almacenamiento documental privado. La interfaz se desarrollará con Astro, componentes React y Tailwind/CSS responsive, manteniendo el criterio de no depender de Bootstrap ni de plantillas completas. El backend se plantea con NestJS sobre Node.js, organizado por módulos, controladores y servicios para autenticación, gestión documental, auditoría e integración con Firma.cl. La base de datos productiva será Supabase PostgreSQL, utilizando TypeORM para el acceso a datos, mientras que los PDFs originales, firmados y recursos asociados se almacenarán en Supabase Storage."
        ),
        "Para reducir el riesgo de dependencia directa del proveedor externo durante la evaluación académica": (
            "Para reducir el riesgo de dependencia directa del proveedor externo durante la evaluación académica, la integración con Firma.cl se implementará mediante una capa de servicio desacoplada dentro del backend NestJS. Esta capa permitirá consumir el servicio real en un ambiente habilitado y, si no existe disponibilidad durante las pruebas, utilizar respuestas controladas que permitan demostrar el flujo principal: carga del documento, selección visual de firma, solicitud de firma, cambio de estado, trazabilidad en Supabase y descarga del resultado."
        ),
        "Desde el punto de vista del sistema, el proyecto utiliza una arquitectura web cliente-servidor.": (
            "Desde el punto de vista del sistema, el proyecto utiliza una arquitectura web cliente-servidor apoyada en servicios gestionados. El frontend Astro/React permite la interacción del usuario mediante formularios, carga de documentos y una vista para ubicar la firma; el backend NestJS procesa reglas de negocio, validaciones, autenticación, almacenamiento, auditoría y comunicación con servicios externos; Supabase PostgreSQL mantiene usuarios, roles, documentos, firmas, estados y auditoría de manera normalizada; y Supabase Storage resguarda los archivos PDF originales y firmados mediante rutas privadas."
        ),
        "Desarrollar un sistema web responsive para la gestión, posicionamiento visual, firma electrónica avanzada": (
            "Desarrollar un sistema web responsive para la gestión, posicionamiento visual, firma electrónica avanzada y descarga de documentos PDF mediante integración con Firma.cl, considerando validación de identidad cuando corresponda al flujo del proveedor, perfiles de usuario diferenciados, base de datos Supabase PostgreSQL, almacenamiento Supabase Storage, trazabilidad de acciones, despliegue en Netlify y formularios funcionales durante el periodo académico de desarrollo del proyecto."
        ),
        "2. Diseñar una arquitectura web compuesta por frontend, backend y base de datos relacional": (
            "2. Diseñar una arquitectura web compuesta por frontend Astro/React, backend NestJS, base de datos Supabase PostgreSQL y almacenamiento Supabase Storage, definiendo módulos, responsabilidades y flujos de comunicación con el proveedor externo de firma electrónica avanzada."
        ),
        "4. Implementar un módulo de documentos que permita cargar archivos PDF": (
            "4. Implementar un módulo de documentos que permita cargar archivos PDF, validar su formato, registrar información relevante en Supabase PostgreSQL, consultar historial y descargar el documento firmado desde Supabase Storage."
        ),
        "6. Diseñar e implementar una base de datos relacional normalizada": (
            "6. Diseñar e implementar una base de datos relacional normalizada en Supabase PostgreSQL con al menos seis tablas, destinada a almacenar usuarios, roles, documentos, firmas, estados, configuraciones visuales de firma, rutas de objetos en Storage y registros de auditoría."
        ),
        "7. Desplegar el sistema en un hosting o servidor web accesible": (
            "7. Desplegar el sistema en Netlify, verificando su funcionamiento responsive, seguridad básica, control de errores, variables de entorno y operación de los formularios principales conectados a Supabase."
        ),
        "El alcance del proyecto se define con el objetivo de construir una solución funcional": (
            "El alcance del proyecto se define con el objetivo de construir una solución funcional, evaluable y pertinente al contexto académico y laboral. La plataforma cubrirá el flujo principal de firma electrónica avanzada de documentos PDF, la gestión de usuarios, la selección visual de la firma, la trazabilidad del proceso, el uso de Supabase para datos y archivos, y el despliegue en Netlify como entorno web accesible."
        ),
        "Registro de metadatos del documento en base de datos.": (
            "Registro de metadatos del documento en Supabase PostgreSQL."
        ),
        "Integración con Firma.cl mediante una capa de servicio backend.": (
            "Integración con Firma.cl mediante una capa de servicio backend NestJS."
        ),
        "Despliegue en hosting o servidor web.": (
            "Despliegue en Netlify con conexión a Supabase."
        ),
        "El análisis de factibilidad permite evaluar si el proyecto puede desarrollarse": (
            "El análisis de factibilidad permite evaluar si el proyecto puede desarrollarse con los recursos técnicos, humanos y operativos disponibles. Para este caso, el proyecto se considera viable porque utiliza tecnologías web conocidas, un alcance delimitado, un modelo de datos manejable, Netlify para hosting, Supabase para base de datos y almacenamiento, y una integración externa encapsulada en el backend NestJS. La principal precisión del alcance es que la plataforma no solo debe solicitar una firma, sino también facilitar la ubicación visual de la firma dentro del PDF."
        ),
        "La planificación considera que algunas tareas pueden superponerse": (
            "La planificación considera que algunas tareas pueden superponerse, especialmente el diseño de interfaz, el modelado de datos, la configuración de Supabase y la preparación del backend NestJS. Sin embargo, las tareas críticas de integración, pruebas y despliegue en Netlify dependen de contar previamente con autenticación, gestión documental, selector visual de firma, modelo de datos funcional y almacenamiento de archivos configurado."
        ),
        "Para el proyecto se selecciona una base de datos relacional, preferentemente PostgreSQL.": (
            "Para el proyecto se selecciona Supabase PostgreSQL como base de datos relacional gestionada. Esta elección se justifica porque el sistema requiere almacenar información estructurada y relacionada entre sí: usuarios, roles, documentos, estados, solicitudes de firma, configuraciones visuales, referencias a objetos almacenados y auditoría. Además, el modelo debe procurar tercera forma normal, evitando duplicidad de información y manteniendo integridad referencial entre las entidades principales."
        ),
        "Se descarta una base de datos NoSQL como opción principal": (
            "Se descarta una base de datos NoSQL como opción principal porque el flujo del sistema depende de relaciones claras y trazables. Por ejemplo, cada documento pertenece a un usuario, posee un estado, puede tener una o más solicitudes de firma, se vincula con objetos privados en Supabase Storage y genera registros de auditoría. Este tipo de estructura se representa de manera más natural mediante tablas relacionales, claves primarias, claves foráneas y restricciones de integridad."
        ),
        "La base de datos se utilizará para almacenar información de control operacional": (
            "La base de datos se utilizará para almacenar información de control operacional, no el contenido binario completo de los documentos PDF. Los archivos originales, firmados y recursos asociados se guardarán en buckets privados de Supabase Storage, mientras que Supabase PostgreSQL registrará sus rutas de objeto, hashes, propietarios, estados, ubicación visible de la firma, fechas y eventos de auditoría. En desarrollo local puede utilizarse almacenamiento temporal o cifrado en carpeta, pero la proyección productiva del stack considera Supabase Storage como repositorio documental."
        ),
        "El siguiente diagrama representa el modelo preliminar de datos del sistema.": (
            "El siguiente diagrama representa el modelo preliminar de datos del sistema sobre Supabase PostgreSQL. Se identifican las entidades principales, sus atributos más relevantes y las relaciones necesarias para mantener trazabilidad entre usuarios, documentos, firmas, configuraciones, estados, referencias a objetos de Supabase Storage y eventos de auditoría."
        ),
        "Las relaciones principales del modelo son: un rol puede estar asociado": (
            "Las relaciones principales del modelo son: un rol puede estar asociado a muchos usuarios; un usuario puede cargar muchos documentos; cada documento posee un estado y referencias a objetos privados en Supabase Storage; un documento puede tener una o más solicitudes o intentos de firma; cada firma puede registrar configuración visual, resultado y errores; y los eventos de auditoría permiten reconstruir acciones relevantes del proceso."
        ),
        "La herramienta de prototipado puede ser Figma, Penpot, Canva o una maqueta HTML inicial.": (
            "La herramienta de prototipado puede ser Figma, Penpot, Canva o una maqueta inicial implementada con Astro/React. Para efectos del diseño preliminar, los wireframes muestran los formularios y pantallas clave exigidas por el proyecto: inicio de sesión, carga de documento, configuración visual de firma, resultado, historial y administración de usuarios. En la implementación final se aplicará Tailwind/CSS responsive con criterios propios de diseño, sin depender de Bootstrap ni de plantillas completas."
        ),
        "Para este proyecto se selecciona una metodología incremental con apoyo de tablero Kanban.": (
            "Para este proyecto se selecciona una metodología incremental con apoyo de tablero Kanban. Esta elección se considera adecuada porque el proyecto será desarrollado de forma individual, con entregas académicas progresivas y con módulos claramente separables: autenticación, gestión de usuarios, carga de documentos, configuración visual de firma, integración externa, auditoría, configuración de Supabase y despliegue en Netlify."
        ),
        "El impacto esperado de esta metodología es positivo": (
            "El impacto esperado de esta metodología es positivo, ya que permite priorizar los elementos evaluables del proyecto, mantener un avance medible y reducir la dependencia de factores externos. Además, al separar la integración de firma en una capa desacoplada y apoyar la infraestructura en Netlify y Supabase, el sistema puede demostrar su flujo principal incluso si el proveedor externo no se encuentra disponible durante una instancia académica."
        ),
        "La plataforma propuesta permite aplicar conocimientos relevantes del área tecnológica": (
            "La plataforma propuesta permite aplicar conocimientos relevantes del área tecnológica, integrando desarrollo web con Astro/React, backend NestJS, base de datos Supabase PostgreSQL, almacenamiento Supabase Storage, seguridad, manejo de archivos, consumo de servicios externos y diseño de experiencia de usuario. Además, el proyecto cuenta con un alcance adecuado para una evaluación académica, ya que incluye perfiles diferenciados, formularios funcionales, modelo relacional, despliegue web en Netlify, trazabilidad de acciones, prototipos y metodología de desarrollo."
        ),
        "En conclusión, el sistema web de gestión, posicionamiento visual": (
            "En conclusión, el sistema web de gestión, posicionamiento visual y firma electrónica avanzada de documentos PDF constituye una propuesta innovadora y útil dentro del área de especialidad. Su desarrollo con Netlify, Supabase, Astro/React y NestJS permitiría mejorar procesos documentales, reducir tiempos, entregar una experiencia más simple a los usuarios y fortalecer la trazabilidad de documentos firmados. Como proyecto de título, permite demostrar competencias técnicas y profesionales directamente vinculadas al desarrollo de soluciones tecnológicas aplicadas al entorno laboral."
        ),
    }

    for fragment, replacement in replacements.items():
        replace_first_containing(doc, fragment, replacement)

    table_updates = {
        (0, 8, 2): "Se registran estados, fechas, usuario y acciones relevantes en Supabase PostgreSQL; el PDF firmado queda en Supabase Storage y el usuario lo descarga desde la plataforma.",
        (1, 7, 1): "La solución considerará control de acceso, auditoría, validación de archivos, resguardo de credenciales de integración y almacenamiento privado de PDFs en Supabase Storage.",
        (2, 7, 2): "El sistema debe registrar nombre, tamaño, fecha de carga, usuario propietario, hash, estado inicial y referencia del objeto almacenado en Supabase Storage.",
        (2, 12, 2): "El sistema debe recibir y almacenar el resultado del proceso de firma en Supabase Storage, asociándolo al documento original.",
        (3, 3, 2): "Claves, tokens y credenciales de integración deben administrarse mediante variables de entorno y nunca quedar expuestas al usuario final.",
        (3, 4, 2): "Los documentos cargados deben almacenarse en buckets privados de Supabase Storage y no ser accesibles públicamente sin autorización.",
        (3, 7, 2): "El sistema debe quedar desplegado en Netlify y accesible para revisión académica.",
        (3, 8, 2): "El código debe organizarse por módulos, separando frontend Astro/React, backend NestJS, servicios, entidades TypeORM, configuración y contratos compartidos.",
        (4, 1, 0): "Hosting",
        (4, 1, 1): "Netlify para publicar la aplicación web y gestionar despliegues, variables de entorno y HTTPS.",
        (4, 2, 1): "NestJS sobre Node.js para controlar sesiones, procesar archivos, coordinar el flujo de firma y comunicarse con servicios externos.",
        (4, 3, 1): "Astro, componentes React, Tailwind/CSS responsive y JavaScript/TypeScript para formularios, navegación, carga de PDF y selector visual de firma.",
        (4, 4, 1): "Supabase PostgreSQL con al menos seis tablas normalizadas y acceso mediante TypeORM.",
        (4, 5, 1): "Supabase Storage para documentos originales, documentos firmados y recursos asociados; almacenamiento local cifrado solo como apoyo de desarrollo.",
        (4, 6, 1): "Capa de servicio NestJS para Firma.cl, con credenciales administradas de forma segura mediante variables de entorno.",
        (4, 7, 1): "Hash de contraseñas, cookies/sesiones seguras, validación de archivos, control por roles y políticas de acceso a Storage.",
        (5, 1, 0): "Hosting productivo",
        (5, 1, 1): "Netlify con despliegues desde repositorio, HTTPS, previews de despliegue y configuración de entorno.",
        (5, 2, 1): "Supabase PostgreSQL con respaldos, roles de acceso, migraciones versionadas y conexión segura desde el backend.",
        (5, 3, 1): "Supabase Storage con buckets privados, políticas de acceso, rutas por documento y reglas de retención.",
        (5, 4, 1): "Funciones programadas, cola externa o mecanismo equivalente para procesar firmas asincrónicas si aumenta el volumen.",
        (5, 5, 1): "Logs de Netlify, registros del backend, métricas de Supabase y alertas por errores de integración.",
        (5, 7, 1): "HTTPS obligatorio, variables de entorno, protección de formularios, limitación de intentos, revisión de dependencias y políticas de acceso en Supabase.",
        (6, 4, 2): "Controlar acceso, usar buckets privados de Supabase Storage, registrar auditoría, aplicar políticas de descarga y definir retención.",
        (7, 6, 1): "Configuración del backend NestJS, Supabase PostgreSQL y Supabase Storage",
        (7, 13, 1): "Despliegue en Netlify y conexión con Supabase",
        (8, 3, 1): "id_documento, id_usuario_propietario, id_estado, nombre_original, bucket_original, object_path_original, bucket_firmado, object_path_firmado, hash_original, hash_firmado, tamano_bytes, fecha_carga",
        (8, 3, 2): "Registra documentos cargados, sus metadatos y las referencias a objetos privados en Supabase Storage.",
        (8, 5, 1): "id_firma, id_documento, id_usuario_firmante, id_estado_firma, pagina, posicion_visual, firma_visible, usa_tsa, fecha_solicitud, fecha_respuesta, mensaje_error, provider_context",
        (8, 8, 1): "id_configuracion, id_usuario, nombre_configuracion, pagina_default, posicion_default, firma_visible_default, usa_qr_default, usa_tsa_default, bucket_default",
        (12, 1, 2): "Estructura inicial del monorepo, rutas, estilos base, variables de entorno, conexión Supabase y despliegue preliminar en Netlify.",
        (12, 3, 2): "Carga de PDF, validación, registro de metadatos, almacenamiento en Supabase Storage e historial.",
        (12, 6, 2): "Publicación en Netlify, conexión a Supabase, revisión responsive y preparación de defensa.",
        (13, 3, 3): "El documento se registra con estado inicial, metadatos y referencia de objeto en Supabase Storage.",
        (13, 6, 3): "El sistema actualiza estado, registra auditoría y habilita descarga del PDF firmado desde Supabase Storage.",
        (13, 10, 3): "Se registran eventos con usuario, fecha, acción, IP, resultado y contexto relevante en Supabase PostgreSQL.",
    }

    for table_idx, row_idx, col_idx in table_updates:
        set_cell_text(
            doc.tables[table_idx].rows[row_idx].cells[col_idx],
            table_updates[(table_idx, row_idx, col_idx)],
        )

    update_references(doc)
    update_anexo(doc)
    doc.save(str(DOC_PATH))


if __name__ == "__main__":
    main()
